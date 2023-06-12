from selenium.webdriver.common.by import By
from selenium import webdriver
import pprint
import time
import pandas as pd
import jieba
from docx import Document
from selenium.common.exceptions import NoSuchElementException
import warnings

from selenium.webdriver.edge.options import Options



def cratch():
    counts = 1
    for link in links:
        print('下载：',link)
        driver.get(link)
        time.sleep(1)
        try:
                txt = driver.find_element(By.XPATH, '//*[@id="Content"]').text
        except NoSuchElementException:
                try:
                    txt = driver.find_element(By.XPATH, '//*[@class="artTxt"]').text
                except NoSuchElementException:
                    txt = ""  # 如果无法定位到日期元素，将txt设置为空字符串或其他默认值

        words.append(txt)
        counts += 1
        print(f"finish {counts} pages of {len(links)}")
        





if __name__ == "__main__":
    # 使用selenium完整的完成以前单纯使用爬取一个步骤完成的任务
    # driver = webdriver.Chrome()
    options = Options()
    options.add_argument('headless')
    options.add_experimental_option('excludeSwitches', ['enable-logging'])
    driver = webdriver.Edge(executable_path=r'd:\edgedriver_win64\msedgedriver.exe',options=options)
    driver.implicitly_wait(30)
    titles,words,dates = [],[],[]
    #进入网站
    with open(r'C:\Users\Scott\CPP\我写的爬虫(23\4\1)\data2.txt', 'a') as file:  # 'a' 表示以追加模式打开文件，如果文件不存在则创建新文件
        links = file.readlines()  # 读取所有行
        links = [i.strip() for i in links]
    with open(r'C:\Users\Scott\CPP\我写的爬虫(23\4\1)\title.txt', 'a') as file:  # 'a' 表示以追加模式打开文件，如果文件不存在则创建新文件
        titles = file.readlines()  # 读取所有行
        titles = [i.strip() for i in links]
    with open(r'C:\Users\Scott\CPP\我写的爬虫(23\4\1)\date.txt', 'a') as file:  # 'a' 表示以追加模式打开文件，如果文件不存在则创建新文件
        dates = file.readlines()  # 读取所有行
        dates = [i.strip() for i in links]
       
        
    cratch()
    driver.close()

    # 创建一个新的Word文档
    doc = Document()

    # 遍历文章列表，添加到Word文档中
    for title, date, content in zip(titles, dates, words):
        # 添加日期
        doc.add_paragraph("日期：" + date)
        doc.add_paragraph("")  # 添加一个空行

        # 添加标题
        doc.add_heading(title, level=1)
        doc.add_paragraph("")  # 添加一个空行

        # 添加文章内容
        doc.add_paragraph(content)
        doc.add_page_break()  # 添加分页符，将每篇文章放在单独的页面上

    # 保存Word文档
    doc.save(r"d:\outpust.docx")