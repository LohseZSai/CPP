from selenium.webdriver.common.by import By
from selenium import webdriver
import pprint
import time
import pandas as pd
import jieba
from docx import Document
from selenium.common.exceptions import NoSuchElementException
import warnings
import pandas as pd
from selenium.webdriver.edge.options import Options

import threading
from concurrent.futures import ThreadPoolExecutor


def cratch():
    """线程任务函数"""
    ct = threading.current_thread()
    print(f"----------{ct.name} 开始----------")

    # 准备webdriver，用于控制浏览器
    options = Options()
    options.add_argument('headless')
    options.add_experimental_option('excludeSwitches', ['enable-logging'])
    driver = webdriver.Edge(executable_path=r'd:\edgedriver_win64\msedgedriver.exe', options=options)
    driver.implicitly_wait(30)

    # 开始下载页面并抓取数据
    counts = 0
    while len(links) > 0:
        link = links.pop(0)
        result = (df[df['urls'] == link]['i'].index)[0]
        print(f'{ct.name} 下载：', link)
        driver.get(link)
        time.sleep(0.1)
        try:
            txt = driver.find_element(By.XPATH, '//*[@id="Content"]').text
        except NoSuchElementException:
            try:
                txt = driver.find_element(By.XPATH, '//*[@class="artTxt"]').text
            except NoSuchElementException:
                txt = f"(maybe images, you can see it on{link})"  # 如果无法定位到日期元素，将txt设置为空字符串或其他默认值

        words.append((result,titles[result],dates[result],txt))
        counts += 1
        print(f"{ct.name} 累计下载 {counts}")
    print(f"----------{ct.name} 结束----------")
    driver.close()


if __name__ == "__main__":
    options = Options()
    options.add_argument('headless')
    options.add_experimental_option('excludeSwitches', ['enable-logging'])

    with open(r'C:\Users\Scott\CPP\我写的爬虫(23\4\1)\data2.txt', 'r') as file:
        links = ([line.strip() for line in file.readlines()])
        urls = [(i, link) for i, link in zip(range(len(links)), links)]
        df = pd.DataFrame(urls, columns=['i', 'urls'])
    with open(r'C:\Users\Scott\CPP\我写的爬虫(23\4\1)\title.txt', 'r') as file:
        titles = [line.strip() for line in file.readlines()]

    with open(r'C:\Users\Scott\CPP\我写的爬虫(23\4\1)\date.txt', 'r') as file:
        dates = [line.strip() for line in file.readlines()]

    words = []
    start = time.time()
    # 使用selenium完整的完成以前单纯使用爬取一个步骤完成的任务
    # driver = webdriver.Chrome()
    # 进入网站

    THREAD_NUM = 30 # 线程数

    # 创建一个包含若干条线程的线程池
    pool = ThreadPoolExecutor(max_workers=THREAD_NUM)
    # 向线程池提交若干个task, urls和new_urls会作为translate()函数的参数
    futures = []
    for i in range(THREAD_NUM):
        future = pool.submit(cratch)
        print(f'提交第{i + 1}个线程任务到线程池')
        # future.add_done_callback(finished)  # 为线程任务注册一个“回调函数”，当该任务成功完成时，程序会自动触发该函数
        futures.append(future)

    # 关闭线程池
    pool.shutdown(wait=True)
    sorted_words = sorted(words, key=lambda x: x[0])
    # 创建一个新的Word文档
    doc = Document()
    for _, title, date, content in sorted_words:
        doc.add_paragraph("日期：" + date)
        doc.add_paragraph("")
        doc.add_heading(title, level=1)
        doc.add_paragraph("")
        doc.add_paragraph(content)
        doc.add_page_break()

    doc.save(r"C:\Users\Scott\Desktop\ffinaloutput.docx")