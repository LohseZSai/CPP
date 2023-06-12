from selenium.webdriver.common.by import By
from selenium import webdriver
import pprint
import time
import pandas as pd
import jieba
from docx import Document
from selenium.common.exceptions import NoSuchElementException

def cratch():
    driver.get('https://newssearch.chinadaily.com.cn/en/search?cond=%7B%22titleMust%22%3A%22energy%22%2C%22fullMust%22%3A%22energy%22%2C%22sort%22%3A%22dp%22%2C%22duplication%22%3A%22on%22%7D&language=en')
    driver.maximize_window()
    time.sleep(3)
    count = 1
    for j in range(50):
        if j == 0:
            nextbutton = driver.find_element(By.XPATH,'/html/body/div[5]/div[2]/div[5]/div[1]/div[2]/span/a[5]')
        else:
            nextbutton1 = driver.find_element(By.XPATH,'/html/body/div[5]/div[2]/div[5]/div[1]/div[2]/span/a[5]')
            driver.execute_script("arguments[0].click();", nextbutton1)
            nextbutton = driver.find_element(By.XPATH,'/html/body/div[5]/div[2]/div[5]/div[1]/div[2]/span/a[6]')
            driver.execute_script("arguments[0].click();", nextbutton)
        for i in range(1,11):
            button =  driver.find_element(By.XPATH,f'/html/body/div[5]/div[2]/div[5]/div[3]/div[{i}]/span/h4/a')  
            title = button.text
            date = driver.find_element(By.XPATH,f'/html/body/div[5]/div[2]/div[5]/div[3]/div[{i}]/span/b[1]').text
            dates.append(date)
            titles.append(title)
            driver.execute_script("arguments[0].click();", button)  
            time.sleep(2)
            allhandle=driver.window_handles
            driver.switch_to.window(allhandle[count])
            count+=1
            try:
                txt = driver.find_element(By.XPATH, '//*[@id="Content"]').text
            except NoSuchElementException:
                try:
                    txt = driver.find_element(By.XPATH, '//*[@class="artTxt"]').text
                except NoSuchElementException:
                    txt = ""  # 如果无法定位到日期元素，将txt设置为空字符串或其他默认值

            words.append(txt)
            time.sleep(1)
            driver.get('https://newssearch.chinadaily.com.cn/en/search?cond=%7B%22titleMust%22%3A%22energy%22%2C%22fullMust%22%3A%22energy%22%2C%22sort%22%3A%22dp%22%2C%22duplication%22%3A%22on%22%7D&language=en')
            time.sleep(2)
            print(f"finish {i} in 50")
if __name__ == "__main__":
    # 使用selenium完整的完成以前单纯使用爬取一个步骤完成的任务

    #进入网站
    titles,words,dates = [],[],[]
    driver = webdriver.Edge()        
    cratch()
    driver.close()

    # 创建一个新的Word文档
    doc = Document()

    # 遍历文章列表，添加到Word文档中
    for title, date, content in zip(titles, dates, words):
        # 添加日期
        doc.add_paragraph("日期：" + date)
        doc.add_paragraph("")  # 添加一个空行

        # 添加标题
        doc.add_heading(title, level=1)
        doc.add_paragraph("")  # 添加一个空行

        # 添加文章内容
        doc.add_paragraph(content)
        doc.add_page_break()  # 添加分页符，将每篇文章放在单独的页面上

    # 保存Word文档
    doc.save(r"C:\Users\Scott\Desktop\output.docx")



