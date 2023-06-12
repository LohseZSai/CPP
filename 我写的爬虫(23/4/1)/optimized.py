from selenium.webdriver.common.by import By
from selenium import webdriver
from selenium.webdriver.edge.options import Options
from selenium.common.exceptions import NoSuchElementException
from docx import Document
import logging
import time

def scrape_links(file_path):
    with open(file_path, 'r') as file:
        links = file.readlines()
        links = list(set(links))
        links = [link.strip() for link in links]
    return links

def scrape_data(links):
    options = Options()
    options.use_chromium = True  # 使用Chromium浏览器内核
    options.add_argument('--headless')  # 使用Headless模式
    driver = webdriver.Edge(options=options)
    time.sleep(0.5)

    titles, words, dates = [], [], []
    counts = 1

    for link in links:
        driver.get(link)

        try:
            txt = driver.find_element(By.XPATH, '//*[@id="Content"]').text
        except NoSuchElementException:
            try:
                txt = driver.find_element(By.XPATH, '//*[@class="artTxt"]').text
            except NoSuchElementException:
                txt = ""

        words.append(txt)
        counts += 1
        print(f"Finished scraping {counts} pages of {len(links)}")

    driver.quit()
    return words

def save_to_word(titles, dates, words, output_path):
    doc = Document()

    for title, date, content in zip(titles, dates, words):
        doc.add_paragraph("日期：" + date)
        doc.add_paragraph("")

        doc.add_heading(title, level=1)
        doc.add_paragraph("")

        doc.add_paragraph(content)
        doc.add_page_break()

    doc.save(output_path)

if __name__ == "__main__":
    file_path = r'C:\Users\Scott\CPP\我写的爬虫(23\4\1)\data.txt'  # 替换为你的文件路径
    links = scrape_links(file_path)

    words = scrape_data(links)

    output_path = r"C:\Users\Scott\Desktop\output.docx"  # 替换为输出的文件路径
    save_to_word([], [], words, output_path)