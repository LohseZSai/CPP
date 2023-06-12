from selenium.webdriver.common.by import By
from selenium import webdriver
import time
import pandas as pd
from docx import Document
from selenium.common.exceptions import NoSuchElementException
import warnings
from selenium.webdriver.edge.options import Options
from concurrent.futures import ThreadPoolExecutor

def scrape(link):
    driver = webdriver.Edge(executable_path=r'd:\edgedriver_win64\msedgedriver.exe', options=options)
    driver.implicitly_wait(30)
    try:
        driver.get(link)
        time.sleep(1)
        try:
            txt = driver.find_element(By.XPATH, '//*[@id="Content"]').text
        except NoSuchElementException:
            try:
                txt = driver.find_element(By.XPATH, '//*[@class="artTxt"]').text
            except NoSuchElementException:
                txt = ""  # 如果无法定位到日期元素，将txt设置为空字符串或其他默认值
        return txt
    finally:
        driver.quit()

if __name__ == "__main__":
    options = Options()
    options.add_argument('headless')
    options.add_experimental_option('excludeSwitches', ['enable-logging'])

    with open(r'C:\Users\Scott\CPP\我写的爬虫(23\4\1)\data2.txt', 'r') as file1:
        links = [line.strip() for line in file1.readlines()]

    with open(r'C:\Users\Scott\CPP\我写的爬虫(23\4\1)\title.txt', 'r') as file2:
        titles = [line.strip() for line in file2.readlines()]

    with open(r'C:\Users\Scott\CPP\我写的爬虫(23\4\1)\date.txt', 'r') as file3:
        dates = [line.strip() for line in file3.readlines()]

    words = []

    with ThreadPoolExecutor(max_workers=8) as executor:
        results = executor.map(scrape, links)
        words = list(results)

    doc = Document()

    for title, date, content in zip(titles, dates, words):
        doc.add_paragraph("日期：" + date)
        doc.add_paragraph("")
        doc.add_heading(title, level=1)
        doc.add_paragraph("")
        doc.add_paragraph(content)
        doc.add_page_break()

    doc.save(r"C:\Users\Scott\Desktop\finaloutput.docx")